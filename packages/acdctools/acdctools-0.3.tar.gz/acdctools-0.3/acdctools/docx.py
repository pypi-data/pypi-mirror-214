import os
import re

try:
    from docx import Document
except ModuleNotFoundError:
    print('='*50)
    print(
        '[ERROR]: acdctools.docx requires the "python-docx" package. '
        'Please install it using the command `pip install python-docx`.'
    )
    print('='*50)
    exit()

def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph

def replace_text_docx(
        docx_path: os.PathLike, new_text: str, text_to_replace: str,
        save_to: os.PathLike=''
    ):
    regex = re.compile(text_to_replace)
    document = Document(docx_path)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_replace_text(paragraph, regex, new_text)

    for p, paragraph in enumerate(document.paragraphs):
        print(paragraph.text)
        paragraph_replace_text(paragraph, regex, new_text)
    
    if not save_to:
        return
    
    document.save(save_to)